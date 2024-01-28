from docx import Document
from docx.shared import RGBColor
import pyperclip


def create_colored_text_docx(text, color):
    # Utworzenie dokumentu
    doc = Document()
    p = doc.add_paragraph()

    # Dodanie kolorowego tekstu
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(color[0], color[1], color[2])

    # Zapisanie dokumentu
    doc_path = 'colored_text.docx'
    doc.save(doc_path)

    return doc_path


def copy_text_from_docx(doc_path):
    # Odczytanie tekstu z dokumentu (bez formatowania)
    doc = Document(doc_path)
    text = doc.paragraphs[0].text

    # Skopiowanie tekstu do schowka
    pyperclip.copy(text)


# Użycie funkcji
color = (255, 0, 0)  # Czerwony kolor
doc_path = create_colored_text_docx("Czerwony tekst", color)
copy_text_from_docx(doc_path)

# Teraz tekst jest w schowku, możesz go wkleić do dokumentu Word
